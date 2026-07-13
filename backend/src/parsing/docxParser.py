"""Defines the DocxParser class for processing DOCX documents."""

from io import BytesIO
from uuid import UUID

from docx import Document

from src.documents.canonical import (
    CanonicalDocument,
    DocumentMetadata,
)
from src.parsing.baseParser import BaseParser


class DocxParser(BaseParser):
    """Office Open XML (.docx) document stream parser."""

    def parse(self, data: bytes, document_id: UUID) -> CanonicalDocument:
        """Parse DOCX bytes into a CanonicalDocument structure."""
        try:
            document = Document(BytesIO(data))
        except Exception as exception:
            raise ValueError(f"Failed to parse DOCX file: {exception}") from exception

        doc_metadata = DocumentMetadata(
            schema_version="1.0.0",
            parser_name="DocxParser",
            parser_version="1.0.0",
        )

        blocks = []
        index = 0
        for paragraph in document.paragraphs:
            text = (paragraph.text or "").strip()
            if not text:
                continue

            block = self.create_default_block(document_id, text, index, page_number=1)
            blocks.append(block)
            index += 1

        return CanonicalDocument(
            document_id=document_id,
            metadata=doc_metadata,
            blocks=blocks,
        )
